2import sys
sys.stdout.reconfigure(encoding='utf-8')
import docx
from docx.oxml.ns import qn

def set_para_text(para, new_text):
    p = para._p
    for r in p.findall(qn('w:r')):
        p.remove(r)
    if new_text:
        para.add_run(new_text)

def set_cell_text(cell, new_text):
    for para in cell.paragraphs:
        p = para._p
        for r in p.findall(qn('w:r')):
            p.remove(r)
    if cell.paragraphs and new_text:
        cell.paragraphs[0].add_run(new_text)
    elif new_text:
        cell.add_paragraph(new_text)

def del_row(table, idx):
    table._tbl.remove(table.rows[idx]._tr)

doc = docx.Document('docs/whatsapp_iot_requirements_v2_en_1.docx')

# 1. Table 1 (metadata): Status
t1 = doc.tables[1]
set_cell_text(t1.rows[2].cells[1], 'Updated — As-Built (reflects deployed implementation)')

# 2. Abstract paragraph
for para in doc.paragraphs:
    if para.text.startswith('Abstract: This document consolidates'):
        set_para_text(para,
            'Abstract: This document consolidates all requirements gathered across the full '
            'design conversation, combined with analysis of the IPX800 V4 hardware manual. '
            'UPDATED: This version reflects the as-built implementation. '
            'Key divergence from the original design: the VPS and IPX800 are on the same LAN '
            '-- no DDNS, no Cloudflare Tunnel, and no router port forwarding are used. '
            'The authoritative as-built reference is docs/documentation.md.')
        break

# 3. Para 83: Network topology description
set_para_text(doc.paragraphs[83],
    'The following topology reflects the as-built deployment: VPS and IPX800 are on the '
    'same LAN. The FastAPI service reaches the IPX800 directly by its local IP address -- '
    'no DDNS, no port forwarding, no router traversal is required.')

# 4. Table 13: Network topology table
t13 = doc.tables[13]
set_cell_text(t13.rows[1].cells[2],
    'Public IP + domain name. Hosts both backend services (Node.js + Python FastAPI). '
    'Reaches IPX800 directly via shared LAN -- no DDNS resolution needed.')
set_cell_text(t13.rows[3].cells[0], 'Home Router')
set_cell_text(t13.rows[3].cells[1], 'Home LAN gateway')
set_cell_text(t13.rows[3].cells[2],
    'VPS and IPX800 share the same LAN. No port forwarding configured. '
    'IPX800 assigned a static LAN IP in its Web UI.')
set_cell_text(t13.rows[4].cells[2],
    'Static LAN IP (e.g. 192.168.1.x, set in IPX800 Web UI). '
    'Reachable directly from VPS: http://<IPX800-LAN-IP>:80 -- no external exposure.')
del_row(t13, 2)  # Remove DuckDNS row

# 5. Data-flow paragraphs 107-110
set_para_text(doc.paragraphs[107],
    '  |  9. GET http://<IPX800-LAN-IP>:80/preset.htm?led1=1&apikey=XXX')
set_para_text(doc.paragraphs[108],
    '  |     (direct LAN -- VPS and IPX800 on same network, no DDNS needed)')
set_para_text(doc.paragraphs[109],
    '  |  10. Read back /status.xml to confirm execution')
set_para_text(doc.paragraphs[110], '')

# 6. Para 143: FastAPI description
set_para_text(doc.paragraphs[143],
    'The Python FastAPI service handles all hardware interaction. It is an internal service '
    'called by Node.js. It reaches the IPX800 directly via its local LAN IP address -- '
    'no DDNS or port forwarding needed since VPS and IPX800 share the same LAN.')

# 7. Table 15: Tech stack
t15 = doc.tables[15]
set_cell_text(t15.rows[8].cells[3],
    'Async HTTP calls to IPX800 via local LAN IP address (direct, no DDNS)')
del_row(t15, 11)  # Port Forwarding row
del_row(t15, 10)  # DDNS row

# 8. Table 17: FastAPI component
t17 = doc.tables[17]
set_cell_text(t17.rows[2].cells[1],
    'Receives command from Node.js, maps to relay number, calls IPX800 via local LAN IP, '
    'applies curtain interlock if needed, returns reply text and success flag')
set_cell_text(t17.rows[3].cells[0], 'GET /health')
set_cell_text(t17.rows[3].cells[1],
    'Pings IPX800 to check reachability via local LAN IP; returns structured JSON health '
    'report with ipx800_host, ipx800_port, and reachability status')
set_cell_text(t17.rows[6].cells[1],
    'Loads IPX800_HOST (<IPX800-LAN-IP>), IPX800_PORT (80), IPX800_APIKEY, relay mapping '
    '(RELAY_LIGHT, RELAY_CURTAIN_UP, RELAY_CURTAIN_DOWN), and log settings from fastapi/.env')
set_cell_text(t17.rows[8].cells[1],
    '3 attempts with 10 s delay -- covers transient LAN errors. After all retries exhausted, '
    'returns success:false to Node.js which sends an error reply to the user.')

# 9. Table 19: Infrastructure prerequisites
t19 = doc.tables[19]
del_row(t19, 3)  # P-03 DuckDNS
del_row(t19, 2)  # P-02 port forwarding
del_row(t19, 1)  # P-01 CG-NAT
set_cell_text(t19.rows[1].cells[0], 'P-01')
set_cell_text(t19.rows[1].cells[1], 'VPS and IPX800 on the same LAN -- static IP on IPX800')
set_cell_text(t19.rows[1].cells[2],
    'Verify both are on the same network segment. Assign IPX800 a static LAN IP in its '
    'Web UI (Network > Parameters > disable DHCP). No DDNS or port forwarding needed.')
set_cell_text(t19.rows[2].cells[0], 'P-02')
set_cell_text(t19.rows[3].cells[0], 'P-03')
set_cell_text(t19.rows[4].cells[0], 'P-04')

# 10. Table 26: Risk register -- remove DDNS/CG-NAT risks
t26 = doc.tables[26]
del_row(t26, 3)  # R3 ISP port blocking
del_row(t26, 2)  # R2 DDNS lag
del_row(t26, 1)  # R1 CG-NAT
for i in range(1, len(t26.rows)):
    set_cell_text(t26.rows[i].cells[0], 'R' + str(i))

# 11. Table 28: Phase 1 steps
t28 = doc.tables[28]
set_cell_text(t28.rows[1].cells[1], 'Verify VPS and IPX800 on the same LAN')
set_cell_text(t28.rows[1].cells[2],
    'Confirm both are on the same network. Ping IPX800 LAN IP from VPS. '
    'Find IPX800 LAN IP in its Web UI (Network > Parameters) or router DHCP table.')
set_cell_text(t28.rows[2].cells[1], 'Assign IPX800 a static LAN IP')
set_cell_text(t28.rows[2].cells[2],
    'IPX800 Web UI: Network > Parameters > disable DHCP > set fixed IP (e.g. 192.168.1.100). '
    'Prevents IP change on DHCP renewal.')
set_cell_text(t28.rows[3].cells[1], 'Set IPX800_HOST and IPX800_APIKEY in fastapi/.env')
set_cell_text(t28.rows[3].cells[2],
    'Edit fastapi/.env: set IPX800_HOST=<IPX800-LAN-IP> (e.g. 192.168.1.100), '
    'IPX800_PORT=80, and IPX800_APIKEY=<your-key>.')
set_cell_text(t28.rows[4].cells[1], 'Set IPX800 to static LAN IP (confirm in Web UI)')
set_cell_text(t28.rows[5].cells[1], 'Configure APIKEY on IPX800')
set_cell_text(t28.rows[5].cells[2],
    'Web UI: Network > API > enable API key protection > set key value > save. '
    'Use same key as IPX800_APIKEY in fastapi/.env.')
set_cell_text(t28.rows[6].cells[1], 'Test relay control from VPS via direct LAN')
set_cell_text(t28.rows[6].cells[2],
    'From VPS: curl "http://<IPX800-LAN-IP>:80/preset.htm?led1=1&apikey=XXX". '
    'Relay 1 (light) should activate. Verify in IPX800 web UI.')
set_cell_text(t28.rows[7].cells[1], 'Test status.xml read from VPS')
set_cell_text(t28.rows[7].cells[2],
    'From VPS: curl "http://<IPX800-LAN-IP>:80/status.xml?apikey=XXX". '
    'Should return XML with <led0>, <led1>, <led2> values.')

# 12. Env config paragraphs
for para in doc.paragraphs:
    t = para.text.strip()
    if t == '# -- IPX800 V4 (accessed via DDNS + port forwarding) -------------':
        set_para_text(para, '# -- IPX800 V4 (accessed via local LAN IP) ---------------------')
    elif '# ── IPX800 V4 (accessed via DDNS' in t:
        set_para_text(para, '# -- IPX800 V4 (accessed via local LAN IP) ---------------------')
    elif t == 'IPX800_HOST=myhome.duckdns.org':
        set_para_text(para, 'IPX800_HOST=<IPX800-LAN-IP>     # e.g. 192.168.1.100')
    elif t == 'IPX800_PORT=8080':
        set_para_text(para, 'IPX800_PORT=80')
    elif '# ── DuckDNS' in t or '# -- DuckDNS' in t:
        set_para_text(para, '')
    elif t == 'DUCKDNS_DOMAIN=myhome':
        set_para_text(para, '')
    elif t == 'DUCKDNS_TOKEN=your-duckdns-token':
        set_para_text(para, '')

doc.save('docs/whatsapp_iot_requirements_v2_en_1.docx')
print('Done. Document saved.')
